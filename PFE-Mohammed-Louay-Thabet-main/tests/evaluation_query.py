# tests/expert_evaluation.py
"""
Expert Evaluation Script — Olive RAG System
============================================
Workflow:
1. Run 20 questions through the local RAG (LLaMA 3.1 8B + TimescaleDB + Qdrant)
2. Save questions + RAG answers to an Excel file
3. Send Excel to the agronomist expert (Dr. specialist in olive)
4. Expert fills in:
   - Their own reference answer
   - Score for the RAG answer (1-5)
   - Specific corrections / missing info
5. Load the filled Excel back → compute evaluation metrics
6. Use expert answers as ground truth for RAGAS and future testing
"""

import os
import sys
import time
import logging
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.agent import ask

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ---------------------------------------------------------------
# 20 Expert-Level Questions
# Covering all major domains of your RAG knowledge base
# ---------------------------------------------------------------

EVALUATION_QUESTIONS = [
    # --- Drought & Climate Stress (5 questions) ---
    {
        "id": "Q01",
        "domain": "Drought Stress",
        "language": "English",
        "question": "What are the physiological mechanisms by which olive trees in Médenine adapt to severe drought conditions, and at what SPI threshold does yield loss become critical?",
    },
    {
        "id": "Q02",
        "domain": "Drought Stress",
        "language": "French",
        "question": "Quels sont les effets du stress hydrique sur la teneur en huile des olives dans la région de Médenine, et comment varient-ils selon le stade phénologique de l'arbre?",
    },
    {
        "id": "Q03",
        "domain": "Drought Stress",
        "language": "English",
        "question": "Based on the climate data from Médenine (1990–2025), which years recorded the most severe drought stress and what was the corresponding impact on olive production?",
    },
    {
        "id": "Q04",
        "domain": "Drought Stress",
        "language": "Arabic",
        "question": "ما هي أكثر سنوات الجفاف حدةً في منطقة مدنين خلال الفترة 1990-2025، وكيف أثرت على إنتاج الزيتون؟",
    },
    {
        "id": "Q05",
        "domain": "Drought Stress",
        "language": "English",
        "question": "How do growing degree days (GDD) correlate with olive yield in Médenine, and what is the optimal GDD range for maximum production?",
    },

    # --- Irrigation & Water Management (3 questions) ---
    {
        "id": "Q06",
        "domain": "Water Management",
        "language": "French",
        "question": "Quelle quantité d'eau consomme un olivier adulte par mois dans les conditions climatiques de Médenine, et quelles sont les meilleures pratiques d'irrigation déficitaire?",
    },
    {
        "id": "Q07",
        "domain": "Water Management",
        "language": "English",
        "question": "What is the relationship between chilling hours and olive flowering in Médenine, and how has climate change affected the chilling hour accumulation since 1990?",
    },
    {
        "id": "Q08",
        "domain": "Water Management",
        "language": "Arabic",
        "question": "ما هي أفضل تقنيات الري بالتنقيط لأشجار الزيتون في المناطق الجافة مثل مدنين، وكيف يمكن تحسين كفاءة استخدام المياه؟",
    },

    # --- Production & Yield Analysis (4 questions) ---
    {
        "id": "Q09",
        "domain": "Production Analysis",
        "language": "English",
        "question": "What is the average annual olive production trend in Médenine from 1990 to 2025, and what climatic factors explain the years with highest and lowest yields?",
    },
    {
        "id": "Q10",
        "domain": "Production Analysis",
        "language": "French",
        "question": "Comment la production oléicole de Médenine se compare-t-elle aux moyennes nationales tunisiennes, et quels facteurs expliquent les écarts observés?",
    },
    {
        "id": "Q11",
        "domain": "Production Analysis",
        "language": "English",
        "question": "What is the alternate bearing pattern observed in Médenine olive orchards, and how can farmers mitigate production variability between on-years and off-years?",
    },
    {
        "id": "Q12",
        "domain": "Production Analysis",
        "language": "Arabic",
        "question": "ما هي العوامل المناخية الرئيسية التي تؤثر على جودة زيت الزيتون في منطقة مدنين، وكيف يمكن تحسين نسبة الزيت في الثمار؟",
    },

    # --- Soil & Agronomy (3 questions) ---
    {
        "id": "Q13",
        "domain": "Soil & Agronomy",
        "language": "English",
        "question": "What soil characteristics are most critical for olive cultivation in the Médenine region, and how should farmers manage soil salinity and pH for optimal yields?",
    },
    {
        "id": "Q14",
        "domain": "Soil & Agronomy",
        "language": "French",
        "question": "Quelles pratiques de taille et de fertilisation sont recommandées pour les oliviers dans les conditions pédoclimatiques de Médenine afin d'optimiser la production?",
    },
    {
        "id": "Q15",
        "domain": "Soil & Agronomy",
        "language": "English",
        "question": "How does wind speed and direction in Médenine affect olive pollination and fruit set, and what protective measures can farmers take during the flowering period?",
    },

    # --- Disease & Pest Management (2 questions) ---
    {
        "id": "Q16",
        "domain": "Disease & Pest",
        "language": "French",
        "question": "Quelles sont les principales maladies et ravageurs affectant les oliviers à Médenine, et comment les conditions climatiques locales influencent-elles leur prévalence?",
    },
    {
        "id": "Q17",
        "domain": "Disease & Pest",
        "language": "English",
        "question": "What are the early warning signs of olive tree decline due to combined drought and heat stress, and what intervention measures are most effective in semi-arid conditions?",
    },

    # --- Sustainability & Climate Adaptation (3 questions) ---
    {
        "id": "Q18",
        "domain": "Sustainability",
        "language": "English",
        "question": "Based on climate projections and historical SPI data for Médenine, what adaptation strategies should olive farmers adopt over the next 10 years to maintain production?",
    },
    {
        "id": "Q19",
        "domain": "Sustainability",
        "language": "French",
        "question": "Comment les technologies innovantes de résistance à la sécheresse, telles que les biostimulants et les variétés tolérantes, peuvent-elles améliorer la durabilité de l'oléiculture à Médenine?",
    },
    {
        "id": "Q20",
        "domain": "Sustainability",
        "language": "Arabic",
        "question": "كيف يمكن لمزارعي الزيتون في مدنين الاستفادة من البيانات المناخية التاريخية لتحسين توقيت عمليات الخدمة الزراعية كالتقليم والتسميد والري؟",
    },
]


# ---------------------------------------------------------------
# Step 1 — Run RAG and collect answers
# ---------------------------------------------------------------

def run_rag_evaluation(backend: str = "llama3.1:8b") -> pd.DataFrame:
    """Run all 20 questions through the RAG and collect answers."""
    
    log.info(f"Starting RAG evaluation with backend: {backend}")
    log.info(f"Total questions: {len(EVALUATION_QUESTIONS)}")
    
    records = []
    
    for i, q in enumerate(EVALUATION_QUESTIONS, 1):
        log.info(f"[{i}/{len(EVALUATION_QUESTIONS)}] {q['id']} — {q['domain']}")
        print(f"\n{'='*60}")
        print(f"[{i}/20] {q['id']} | {q['domain']} | {q['language']}")
        print(f"Q: {q['question'][:100]}...")
        
        start = time.time()
        try:
            result = ask(
                question=q["question"],
                backend=backend,
                use_db=True,
            )
            rag_answer      = result.get("answer", "ERROR: No answer returned")
            response_time   = result.get("response_time", round(time.time() - start, 2))
            retrieved_chunks = _format_chunks(result.get("context", {}))
            
        except Exception as e:
            log.error(f"Error on {q['id']}: {e}")
            rag_answer       = f"ERROR: {str(e)}"
            response_time    = round(time.time() - start, 2)
            retrieved_chunks = ""
        
        print(f"✓ Answer received in {response_time}s")
        print(f"  Preview: {rag_answer[:150]}...")
        
        records.append({
            "question_id": q["id"],
            "domain":      q["domain"],
            "language":    q["language"],
            "question":    q["question"],
            "rag_answer":  rag_answer,
        })
        
        # Small delay to avoid overwhelming Ollama
        time.sleep(1)
    
    return pd.DataFrame(records)


def _format_chunks(context: dict) -> str:
    """Format retrieved chunks for display in Excel."""
    sql = context.get("sql_context", "")
    vec = context.get("vector_context", "")
    result = ""
    if sql:
        result += f"[SQL DATA]\n{sql[:500]}\n\n"
    if vec:
        result += f"[VECTOR DOCS]\n{vec[:800]}"
    return result.strip()


# ---------------------------------------------------------------
# Step 2 — Export to Word document for expert review
# ---------------------------------------------------------------

def export_for_expert(df: pd.DataFrame, output_path: str):
    """Export Q&A pairs to a clean Word document — domain, language, question, answer only."""

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────
    title = doc.add_heading("Olive RAG System — Expert Evaluation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph("RAG Answers for Expert Review — Médenine Olive Agronomy")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()  # spacer

    # ── One entry per question ─────────────────────────────────
    for _, row in df.iterrows():
        # Question header  e.g.  "Q01 — Drought Stress"
        heading = doc.add_heading(
            f"{row['question_id']} — {row['domain']}",
            level=1
        )
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        # Metadata line: Language
        meta = doc.add_paragraph()
        meta.add_run("Language: ").bold = True
        meta.add_run(str(row["language"]))
        meta.runs[0].font.size = Pt(10)

        # Question label + text
        q_label = doc.add_paragraph()
        q_label.add_run("Question:  ").bold = True
        q_label.runs[0].font.size = Pt(11)

        q_text = doc.add_paragraph(str(row["question"]))
        q_text.runs[0].font.size   = Pt(11)
        q_text.runs[0].font.italic = True

        # Answer label + text
        a_label = doc.add_paragraph()
        a_label.add_run("RAG Answer:  ").bold = True
        a_label.runs[0].font.size = Pt(11)

        answer_text = str(row.get("rag_answer", "No answer generated."))
        a_text = doc.add_paragraph(answer_text)
        a_text.runs[0].font.size = Pt(11)

        # Divider between entries
        doc.add_paragraph("─" * 80)

    doc.save(output_path)
    log.info(f"Document saved to: {output_path}")
    print(f"\n✅ Word document ready: {output_path}")


# ---------------------------------------------------------------
# Step 3 — Load expert-filled Excel and compute metrics
# ---------------------------------------------------------------

def load_expert_results(filled_excel_path: str) -> dict:
    """
    Load the Excel filled by the expert and compute evaluation metrics.
    Call this AFTER the expert returns the filled file.
    """
    
    df = pd.read_excel(filled_excel_path, sheet_name="Expert_Review")
    
    # Filter rows where expert has filled in scores
    reviewed = df[df["expert_score"].notna() & (df["expert_score"] != "")]
    
    if reviewed.empty:
        print("❌ No expert scores found yet. Expert has not filled in the file.")
        return {}
    
    scores = pd.to_numeric(reviewed["expert_score"], errors="coerce").dropna()
    
    score_labels = {
        5: "Excellent",
        4: "Good",
        3: "Acceptable",
        2: "Poor",
        1: "Very Poor"
    }
    
    metrics = {
        "total_reviewed":        len(reviewed),
        "average_score":         round(scores.mean(), 3),
        "score_distribution":    scores.value_counts().sort_index().to_dict(),
        "pct_acceptable_above":  round((scores >= 3).sum() / len(scores) * 100, 1),
        "pct_good_above":        round((scores >= 4).sum() / len(scores) * 100, 1),
        "pct_excellent":         round((scores == 5).sum() / len(scores) * 100, 1),
        "has_hallucinations":    reviewed["expert_hallucinations"].notna().sum(),
        "by_domain":             reviewed.groupby("domain")["expert_score"]
                                    .apply(lambda x: round(pd.to_numeric(x, errors='coerce').mean(), 2))
                                    .to_dict(),
        "by_language":           reviewed.groupby("language")["expert_score"]
                                    .apply(lambda x: round(pd.to_numeric(x, errors='coerce').mean(), 2))
                                    .to_dict(),
    }
    
    # Print report
    print("\n" + "="*60)
    print("EXPERT EVALUATION RESULTS")
    print("="*60)
    print(f"Questions reviewed : {metrics['total_reviewed']}/20")
    print(f"Average score      : {metrics['average_score']}/5.0")
    print(f"Acceptable+ (≥3)   : {metrics['pct_acceptable_above']}%")
    print(f"Good+ (≥4)         : {metrics['pct_good_above']}%")
    print(f"Excellent (5)      : {metrics['pct_excellent']}%")
    print(f"Hallucinations found: {metrics['has_hallucinations']} answers")
    
    print("\nBy Domain:")
    for domain, score in metrics["by_domain"].items():
        bar = "█" * int(score * 2)
        print(f"  {domain:<25} {score:.2f}/5.0  {bar}")
    
    print("\nBy Language:")
    for lang, score in metrics["by_language"].items():
        bar = "█" * int(score * 2)
        print(f"  {lang:<10} {score:.2f}/5.0  {bar}")
    
    print("\nScore Distribution:")
    for score_val, count in sorted(metrics["score_distribution"].items()):
        label = score_labels.get(int(score_val), "")
        bar   = "█" * count
        print(f"  {int(score_val)} ({label:<12}): {bar} ({count})")
    
    # Save ground truth — expert reference answers for future RAGAS evaluation
    ground_truth = reviewed[["question_id", "question", "expert_reference_answer"]].copy()
    ground_truth = ground_truth[ground_truth["expert_reference_answer"].notna()]
    ground_truth.to_csv("tests/expert_ground_truth.csv", index=False)
    print(f"\n💾 Ground truth saved to: tests/expert_ground_truth.csv")
    print(f"   ({len(ground_truth)} reference answers ready for RAGAS evaluation)")
    
    return metrics


# ---------------------------------------------------------------
# Main
# ---------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Expert Evaluation Script")
    parser.add_argument(
        "--mode",
        choices=["run", "export", "analyze"],
        default="run",
        help=(
            "run     → run RAG on 20 questions and export Excel\n"
            "analyze → load filled Excel and compute metrics"
        )
    )
    parser.add_argument(
        "--backend",
        default="llama3.1:8b",
        help="LLM backend to use (default: llama3.1:8b)"
    )
    parser.add_argument(
        "--input",
        default=None,
        help="Path to expert-filled Excel file (for --mode analyze)"
    )
    args = parser.parse_args()
    
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = f"tests/expert_evaluation_{timestamp}.docx"
    
    if args.mode in ("run", "export"):
        print("\n" + "="*60)
        print("OLIVE RAG — EXPERT EVALUATION")
        print("="*60)
        print(f"Backend : {args.backend}")
        print(f"Output  : {output_path}")
        print(f"Questions: {len(EVALUATION_QUESTIONS)}")
        print("="*60)
        
        # Run RAG
        df = run_rag_evaluation(backend=args.backend)
        
        # Export to Excel
        export_for_expert(df, output_path)
        
        print("\n📋 NEXT STEPS:")
        print(f"  1. Send '{output_path}' to your agronomist expert")
        print(f"  2. Ask them to review each answer and provide their reference answers")
        print(f"  3. Once returned, run:")
        print(f"     python tests/expert_evaluation.py --mode analyze --input <filled_file.xlsx>")
    
    elif args.mode == "analyze":
        if not args.input:
            print("❌ Please provide the filled Excel path: --input path/to/filled_file.xlsx")
            sys.exit(1)
        load_expert_results(args.input)
        